#!/usr/bin/env python3
"""
VulnHunter - Core Vulnerability Detection Engine
Advanced OWASP Top 10 vulnerability scanning for web applications.
"""

import requests
from datetime import datetime
from urllib.parse import urlparse
import re
import time
import random
from bs4 import BeautifulSoup

VERSION = "2.1.0"

class OWASPTop10Scanner:
    def __init__(self, target_url, max_depth=2, delay=1.0):
        self.target_url = target_url.rstrip('/')
        self.max_depth = max_depth
        self.delay = delay
        self.visited_urls = set()
        self.vulnerabilities = []
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
        
    def get_owasp_top10_coverage(self):
        """Returns comprehensive OWASP Top 10 2021 coverage information with CVSS details"""
        return {
            'A01': {
                'name': 'Broken Access Control',
                'description': 'Restrictions on what authenticated users can do are often not properly enforced',
                'tests': ['Directory Traversal', 'Privilege Escalation', 'Force Browsing'],
                'severity_range': 'Critical to Medium',
                'cvss_range': '9.1 - 5.3',
                'impact': 'Unauthorized access to data/functionality, privilege escalation'
            },
            'A02': {
                'name': 'Cryptographic Failures',
                'description': 'Failures related to cryptography which often leads to sensitive data exposure',
                'tests': ['Weak SSL/TLS', 'Insecure Protocols', 'Weak Encryption'],
                'severity_range': 'High to Medium',
                'cvss_range': '7.5 - 5.9',
                'impact': 'Data exposure, man-in-the-middle attacks, credential theft'
            },
            'A03': {
                'name': 'Injection',
                'description': 'Application vulnerable to injection flaws like SQL, NoSQL, OS, and LDAP injection',
                'tests': ['SQL Injection', 'XSS (Reflected/Stored)', 'Command Injection'],
                'severity_range': 'Critical',
                'cvss_range': '9.8 - 8.8',
                'impact': 'Data breach, unauthorized system access, malicious code execution'
            },
            'A04': {
                'name': 'Insecure Design',
                'description': 'Risks related to design and architectural flaws',
                'tests': ['Missing Security Controls', 'Insecure Design Patterns'],
                'severity_range': 'High to Medium',
                'cvss_range': '8.1 - 6.5',
                'impact': 'Business logic bypass, architectural vulnerabilities'
            },
            'A05': {
                'name': 'Security Misconfiguration',
                'description': 'Missing appropriate security hardening across the application stack',
                'tests': ['Missing Security Headers', 'Default Configurations', 'Verbose Error Messages'],
                'severity_range': 'Medium to Low',
                'cvss_range': '6.1 - 4.3',
                'impact': 'Information disclosure, clickjacking, cache poisoning'
            },
            'A06': {
                'name': 'Vulnerable and Outdated Components',
                'description': 'Using components with known vulnerabilities',
                'tests': ['Outdated Libraries', 'Vulnerable Dependencies', 'Unpatched Software'],
                'severity_range': 'Critical to Low',
                'cvss_range': '9.0 - 3.9',
                'impact': 'Remote code execution, data breach, system compromise'
            },
            'A07': {
                'name': 'Identification and Authentication Failures',
                'description': 'Confirmation of user identity, authentication, and session management',
                'tests': ['Weak Passwords', 'Session Management', 'Authentication Bypass'],
                'severity_range': 'High to Medium',
                'cvss_range': '8.1 - 6.5',
                'impact': 'Account takeover, unauthorized access, session hijacking'
            },
            'A08': {
                'name': 'Software and Data Integrity Failures',
                'description': 'Code and infrastructure that do not protect against integrity violations',
                'tests': ['Insecure Deserialization', 'CI/CD Pipeline Issues', 'Auto-Update Issues'],
                'severity_range': 'High to Medium',
                'cvss_range': '8.1 - 5.6',
                'impact': 'Remote code execution, data tampering, supply chain attacks'
            },
            'A09': {
                'name': 'Security Logging and Monitoring Failures',
                'description': 'Insufficient logging and monitoring coupled with missing incident response',
                'tests': ['Missing Logging', 'Inadequate Monitoring', 'Poor Incident Response'],
                'severity_range': 'Medium to Low',
                'cvss_range': '6.5 - 4.0',
                'impact': 'Delayed breach detection, forensic difficulties, compliance issues'
            },
            'A10': {
                'name': 'Server-Side Request Forgery (SSRF)',
                'description': 'SSRF flaws occur when a web app fetches remote resources without validating the URL',
                'tests': ['Internal Resource Access', 'Cloud Metadata Access', 'Port Scanning'],
                'severity_range': 'Critical to Medium',
                'cvss_range': '9.1 - 6.5',
                'impact': 'Internal network access, cloud metadata exposure, service enumeration'
            }
        }
    
    def crawl_website(self):
        """Basic website crawling to discover URLs"""
        try:
            response = self.session.get(self.target_url, timeout=10)
            self.visited_urls.add(self.target_url)
            
            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'html.parser')
                # Find links
                for link in soup.find_all('a', href=True):
                    href = link['href']
                    if href.startswith('/'):
                        full_url = self.target_url + href
                    elif href.startswith(self.target_url):
                        full_url = href
                    else:
                        continue
                    
                    if len(self.visited_urls) < 10:  # Limit crawling
                        self.visited_urls.add(full_url)
        except:
            pass
    
    def scan_a01_broken_access_control(self):
        """A01:2021 – Broken Access Control"""
        vulnerabilities = []
        
        # Directory traversal payloads
        traversal_payloads = [
            '../../../etc/passwd',
            '..\\..\\..\\windows\\system32\\drivers\\etc\\hosts',
            '../../../boot.ini',
            '../../../../../../../../etc/passwd',
            '..%2F..%2F..%2Fetc%2Fpasswd',
            '....//....//....//etc/passwd'
        ]
        
        # Test directory traversal
        for url in list(self.visited_urls)[:5]:
            for payload in traversal_payloads:
                try:
                    test_url = f"{url}?file={payload}"
                    response = self.session.get(test_url, timeout=5)
                    
                    if any(indicator in response.text.lower() for indicator in 
                          ['root:', 'bin:', '[fonts]', 'windows', '/bin/bash', 'daemon:']):
                        vulnerabilities.append({
                            'type': 'Directory Traversal',
                            'severity': 'Critical',
                            'cvss_score': 9.1,
                            'cvss_vector': 'CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N',
                            'category': 'A01',
                            'owasp_category': 'A01:2021 - Broken Access Control',
                            'url': test_url,
                            'parameter': 'file',
                            'payload': payload,
                            'description': 'Critical directory traversal vulnerability allowing access to system files. This violates the principle of least privilege and can lead to complete system compromise.',
                            'impact': 'Unauthorized file system access, potential credential exposure, system information disclosure',
                            'remediation': 'Implement proper input validation, use whitelist approach, avoid direct file path manipulation'
                        })
                    
                    time.sleep(self.delay)
                except:
                    continue
        
        # Test for common admin panels
        admin_paths = ['/admin', '/admin.php', '/administrator', '/wp-admin', '/management', '/control-panel']
        for path in admin_paths:
            try:
                admin_url = self.target_url + path
                response = self.session.get(admin_url, timeout=5)
                if response.status_code == 200 and 'login' in response.text.lower():
                    vulnerabilities.append({
                        'type': 'Exposed Admin Panel',
                        'severity': 'High',
                        'cvss_score': 7.5,
                        'cvss_vector': 'CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N',
                        'category': 'A01',
                        'owasp_category': 'A01:2021 - Broken Access Control',
                        'url': admin_url,
                        'description': 'Administrative interface is publicly accessible without proper access controls',
                        'impact': 'Unauthorized access to administrative functions, potential system compromise',
                        'remediation': 'Implement proper authentication and IP-based access controls for admin panels'
                    })
                time.sleep(self.delay)
            except:
                continue
        
        return vulnerabilities

    def scan_a02_cryptographic_failures(self):
        """A02:2021 – Cryptographic Failures"""
        vulnerabilities = []
        
        try:
            # Check if site uses HTTPS
            if not self.target_url.startswith('https://'):
                vulnerabilities.append({
                    'type': 'Insecure Protocol',
                    'severity': 'High',
                    'cvss_score': 7.4,
                    'cvss_vector': 'CVSS:3.1/AV:N/AC:H/PR:N/UI:N/S:U/C:H/I:H/A:N',
                    'category': 'A02',
                    'owasp_category': 'A02:2021 - Cryptographic Failures',
                    'url': self.target_url,
                    'description': 'Website uses insecure HTTP protocol instead of HTTPS, exposing data to interception',
                    'impact': 'Data interception, man-in-the-middle attacks, credential theft',
                    'remediation': 'Implement HTTPS with proper SSL/TLS certificates and redirect HTTP to HTTPS'
                })
            
            # Test for mixed content
            response = self.session.get(self.target_url, timeout=10)
            if 'http://' in response.text and self.target_url.startswith('https://'):
                vulnerabilities.append({
                    'type': 'Mixed Content',
                    'severity': 'Medium',
                    'cvss_score': 6.1,
                    'cvss_vector': 'CVSS:3.1/AV:N/AC:L/PR:N/UI:R/S:C/C:L/I:L/A:N',
                    'category': 'A02',
                    'owasp_category': 'A02:2021 - Cryptographic Failures',
                    'url': self.target_url,
                    'description': 'HTTPS page contains insecure HTTP resources, creating mixed content vulnerabilities',
                    'impact': 'Partial encryption bypass, potential data interception of mixed resources',
                    'remediation': 'Ensure all resources (images, scripts, stylesheets) are loaded over HTTPS'
                })
        except:
            pass
        
        return vulnerabilities

    def scan_a03_injection(self):
        """A03:2021 – Injection"""
        vulnerabilities = []
        
        # SQL Injection payloads
        sql_payloads = [
            "' OR '1'='1",
            "'; DROP TABLE users; --",
            "' UNION SELECT NULL, version(), NULL --",
            "admin'--",
            "' OR 1=1 --",
            "' OR 'a'='a",
            "1' AND SLEEP(5) --"
        ]
        
        # XSS payloads
        xss_payloads = [
            "<script>alert('XSS')</script>",
            "javascript:alert('XSS')",
            "<img src=x onerror=alert('XSS')>",
            "';alert('XSS');//",
            "<svg onload=alert('XSS')>",
            "'\"><script>alert('XSS')</script>"
        ]
        
        # Test forms for injection vulnerabilities
        try:
            response = self.session.get(self.target_url, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            forms = soup.find_all('form')
            for form in forms:
                form_url = form.get('action', self.target_url)
                if not form_url.startswith('http'):
                    form_url = self.target_url + form_url if form_url.startswith('/') else self.target_url + '/' + form_url
                
                inputs = form.find_all(['input', 'textarea'])
                params = {}
                
                for input_field in inputs:
                    field_name = input_field.get('name', 'test')
                    field_type = input_field.get('type', 'text')
                    
                    if field_type not in ['submit', 'button', 'reset', 'file']:
                        params[field_name] = 'test'
                
                # Test SQL injection
                for param in params:
                    for payload in sql_payloads[:3]:  # Test first 3 payloads
                        test_params = params.copy()
                        test_params[param] = payload
                        
                        try:
                            if form.get('method', '').lower() == 'post':
                                test_response = self.session.post(form_url, data=test_params, timeout=5)
                            else:
                                test_response = self.session.get(form_url, params=test_params, timeout=5)
                            
                            # Check for SQL error messages
                            sql_errors = [
                                'sql syntax', 'mysql_fetch', 'warning: mysql',
                                'valid mysql result', 'postgresql', 'warning: pg_',
                                'valid postgresql result', 'oracle', 'quoted string not properly terminated',
                                'sqlite_exception', 'sqlite3', 'microsoft ole db provider for odbc drivers'
                            ]
                            
                            error_found = any(error in test_response.text.lower() for error in sql_errors)
                            if error_found:
                                vulnerabilities.append({
                                    'type': 'SQL Injection',
                                    'severity': 'Critical',
                                    'cvss_score': 9.8,
                                    'cvss_vector': 'CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H',
                                    'category': 'A03',
                                    'owasp_category': 'A03:2021 - Injection',
                                    'url': form_url,
                                    'parameter': param,
                                    'payload': payload,
                                    'description': 'Critical SQL injection vulnerability - database errors exposed. Allows arbitrary database queries and potential full database compromise.',
                                    'impact': 'Complete database compromise, data breach, unauthorized data modification, potential system takeover',
                                    'remediation': 'Use parameterized queries/prepared statements, implement proper input validation, apply least privilege database access'
                                })
                            
                            # Check for timing-based SQL injection
                            elif 'SLEEP' in payload.upper():
                                start_time = time.time()
                                response = self.session.get(form_url, params=test_params, timeout=10)
                                response_time = time.time() - start_time
                                
                                if response_time > 4:  # 5 second sleep should take ~5 seconds
                                    vulnerabilities.append({
                                        'type': 'Time-based SQL Injection',
                                        'severity': 'Critical',
                                        'cvss_score': 9.8,
                                        'cvss_vector': 'CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H',
                                        'category': 'A03',
                                        'owasp_category': 'A03:2021 - Injection',
                                        'url': form_url,
                                        'parameter': param,
                                        'payload': payload,
                                        'description': f'Time-based SQL injection detected (response time: {response_time:.2f}s)',
                                        'impact': 'Complete database compromise, data exfiltration, potential system access',
                                        'remediation': 'Use parameterized queries, implement proper input validation'
                                    })
                            
                            time.sleep(self.delay)
                        except:
                            continue
                
                # Test XSS
                for param in params:
                    for payload in xss_payloads[:3]:  # Test first 3 payloads
                        test_params = params.copy()
                        test_params[param] = payload
                        
                        try:
                            if form.get('method', '').lower() == 'post':
                                test_response = self.session.post(form_url, data=test_params, timeout=5)
                            else:
                                test_response = self.session.get(form_url, params=test_params, timeout=5)
                            
                            # Check if payload is reflected in response
                            if payload in test_response.text:
                                vulnerabilities.append({
                                    'type': 'Cross-Site Scripting (XSS)',
                                    'severity': 'Critical',
                                    'cvss_score': 8.8,
                                    'cvss_vector': 'CVSS:3.1/AV:N/AC:L/PR:N/UI:R/S:U/C:H/I:H/A:H',
                                    'category': 'A03',
                                    'owasp_category': 'A03:2021 - Injection',
                                    'url': form_url,
                                    'parameter': param,
                                    'payload': payload,
                                    'description': 'Critical XSS vulnerability - unfiltered script execution allows arbitrary JavaScript execution in victim browsers',
                                    'impact': 'Session hijacking, credential theft, malicious redirection, defacement',
                                    'remediation': 'Implement proper output encoding, Content Security Policy, input validation'
                                })
                            
                            time.sleep(self.delay)
                        except:
                            continue
        
        except:
            pass
        
        return vulnerabilities

    def scan_a04_insecure_design(self):
        """A04:2021 – Insecure Design"""
        vulnerabilities = []
        
        try:
            # Check for common insecure design patterns
            response = self.session.get(self.target_url, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Check for forms without CSRF protection
            forms = soup.find_all('form')
            for form in forms:
                csrf_token = form.find('input', {'name': re.compile(r'csrf|token', re.I)})
                if not csrf_token and form.get('method', '').upper() == 'POST':
                    vulnerabilities.append({
                        'type': 'Missing CSRF Protection',
                        'severity': 'High',
                        'cvss_score': 8.1,
                        'cvss_vector': 'CVSS:3.1/AV:N/AC:L/PR:N/UI:R/S:U/C:H/I:H/A:N',
                        'category': 'A04',
                        'owasp_category': 'A04:2021 - Insecure Design',
                        'url': self.target_url,
                        'description': 'Form lacks CSRF token protection, allowing cross-site request forgery attacks',
                        'form_action': form.get('action', 'Not specified'),
                        'impact': 'Unauthorized actions on behalf of authenticated users, account compromise',
                        'remediation': 'Implement CSRF tokens in all state-changing forms and validate them server-side'
                    })
            
            # Check for password fields without proper constraints
            password_fields = soup.find_all('input', {'type': 'password'})
            for field in password_fields:
                if not field.get('minlength') and not field.get('pattern'):
                    vulnerabilities.append({
                        'type': 'Weak Password Policy',
                        'severity': 'Medium',
                        'cvss_score': 5.3,
                        'cvss_vector': 'CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:L/I:N/A:N',
                        'category': 'A04',
                        'owasp_category': 'A04:2021 - Insecure Design',
                        'url': self.target_url,
                        'description': 'Password field lacks complexity requirements, enabling weak password usage',
                        'impact': 'Increased risk of brute force attacks, weak credential creation',
                        'remediation': 'Implement strong password policies with minimum length, complexity requirements'
                    })
                    
            # Check for autocomplete on sensitive fields
            sensitive_fields = soup.find_all('input', {'type': ['password', 'email', 'tel']})
            for field in sensitive_fields:
                if field.get('autocomplete') != 'off':
                    vulnerabilities.append({
                        'type': 'Autocomplete Enabled on Sensitive Field',
                        'severity': 'Low',
                        'cvss_score': 3.1,
                        'cvss_vector': 'CVSS:3.1/AV:P/AC:H/PR:N/UI:N/S:U/C:L/I:N/A:N',
                        'category': 'A04',
                        'owasp_category': 'A04:2021 - Insecure Design',
                        'url': self.target_url,
                        'field_type': field.get('type'),
                        'description': f'Autocomplete enabled on sensitive {field.get("type")} field',
                        'impact': 'Potential credential storage in browser cache',
                        'remediation': 'Set autocomplete="off" on sensitive form fields'
                    })
        
        except:
            pass
        
        return vulnerabilities

    def scan_a05_security_misconfiguration(self):
        """A05:2021 – Security Misconfiguration"""
        vulnerabilities = []
        
        try:
            response = self.session.get(self.target_url, timeout=10)
            headers = response.headers
            
            # Check for missing security headers
            security_headers = {
                'X-Content-Type-Options': 'Prevents MIME-sniffing attacks',
                'X-Frame-Options': 'Prevents clickjacking attacks',
                'X-XSS-Protection': 'Enables XSS filtering in browsers',
                'Strict-Transport-Security': 'Enforces HTTPS connections',
                'Content-Security-Policy': 'Prevents XSS and data injection attacks',
                'Referrer-Policy': 'Controls referrer information leakage'
            }
            
            for header, description in security_headers.items():
                if header not in headers:
                    cvss_score = 6.1 if header in ['Content-Security-Policy', 'X-Frame-Options'] else 4.3
                    vulnerabilities.append({
                        'type': 'Missing Security Header',
                        'severity': 'Medium' if cvss_score > 5.0 else 'Low',
                        'cvss_score': cvss_score,
                        'cvss_vector': 'CVSS:3.1/AV:N/AC:L/PR:N/UI:R/S:C/C:L/I:L/A:N' if cvss_score > 5.0 else 'CVSS:3.1/AV:N/AC:H/PR:N/UI:N/S:U/C:L/I:N/A:N',
                        'category': 'A05',
                        'owasp_category': 'A05:2021 - Security Misconfiguration',
                        'url': self.target_url,
                        'header': header,
                        'description': f'{description}. This security misconfiguration can lead to various client-side attacks.',
                        'impact': 'Clickjacking, XSS attacks, information disclosure, cache poisoning',
                        'remediation': f'Add {header} header with appropriate values to HTTP responses'
                    })
            
            # Check for server information disclosure
            server_header = headers.get('Server', '')
            if server_header and any(tech in server_header.lower() for tech in ['apache', 'nginx', 'iis', 'tomcat']):
                vulnerabilities.append({
                    'type': 'Server Information Disclosure',
                    'severity': 'Low',
                    'cvss_score': 3.7,
                    'cvss_vector': 'CVSS:3.1/AV:N/AC:H/PR:N/UI:N/S:U/C:L/I:N/A:N',
                    'category': 'A05',
                    'owasp_category': 'A05:2021 - Security Misconfiguration',
                    'url': self.target_url,
                    'description': f'Server header reveals technology: {server_header}',
                    'impact': 'Information leakage for reconnaissance, attack surface identification',
                    'remediation': 'Remove or obfuscate server version information in HTTP headers'
                })
        
        except:
            pass
        
        return vulnerabilities

    def scan_a06_vulnerable_components(self):
        """A06:2021 – Vulnerable and Outdated Components"""
        vulnerabilities = []
        
        try:
            response = self.session.get(self.target_url, timeout=10)
            
            # Check for common vulnerable frameworks/libraries in HTML
            vulnerable_patterns = {
                r'jquery[/-](\d+\.\d+\.\d+)': ('jQuery', '3.0.0'),
                r'bootstrap[/-](\d+\.\d+\.\d+)': ('Bootstrap', '4.0.0'),
                r'angular[/-](\d+\.\d+\.\d+)': ('Angular', '8.0.0'),
                r'react[/-](\d+\.\d+\.\d+)': ('React', '16.0.0')
            }
            
            for pattern, (lib_name, min_version) in vulnerable_patterns.items():
                matches = re.findall(pattern, response.text, re.IGNORECASE)
                for version in matches:
                    # Simple version comparison (not fully accurate but indicative)
                    version_parts = [int(x) for x in version.split('.')]
                    min_parts = [int(x) for x in min_version.split('.')]
                    
                    if version_parts < min_parts:
                        vulnerabilities.append({
                            'type': 'Outdated JavaScript Library',
                            'severity': 'Medium',
                            'cvss_score': 6.1,
                            'cvss_vector': 'CVSS:3.1/AV:N/AC:L/PR:N/UI:R/S:C/C:L/I:L/A:N',
                            'category': 'A06',
                            'owasp_category': 'A06:2021 - Vulnerable and Outdated Components',
                            'url': self.target_url,
                            'library': lib_name,
                            'version': version,
                            'description': f'Outdated {lib_name} version {version} detected. May contain known vulnerabilities.',
                            'impact': 'Potential XSS, CSRF, or other client-side vulnerabilities',
                            'remediation': f'Update {lib_name} to the latest stable version'
                        })
        
        except Exception as e:
            pass
            
        # Check for common CMS indicators and versions
        cms_indicators = {
            'wp-content': 'WordPress',
            'sites/default': 'Drupal', 
            'administrator/index.php': 'Joomla',
            'typo3': 'TYPO3',
            'concrete5': 'Concrete5',
            'magento': 'Magento',
            'prestashop': 'PrestaShop'
        }
        
        try:
            response = self.session.get(self.target_url, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Check for version disclosure in meta tags
            meta_generator = soup.find('meta', {'name': 'generator'})
            if meta_generator:
                content = meta_generator.get('content', '')
                vulnerabilities.append({
                    'type': 'Version Disclosure in Meta Tag',
                'severity': 'Low',
                'cvss_score': 3.7,
                'cvss_vector': 'CVSS:3.1/AV:N/AC:H/PR:N/UI:N/S:U/C:L/I:N/A:N',
                'category': 'A06',
                'owasp_category': 'A06:2021 - Vulnerable and Outdated Components',
                'url': self.target_url,
                'version_info': content,
                'description': f'Version information disclosed: {content}',
                'impact': 'Technology stack disclosure aids targeted attacks',
                    'remediation': 'Remove or obfuscate version information from meta tags'
                })
            
            for indicator, cms_name in cms_indicators.items():
                if indicator in response.text.lower():
                    vulnerabilities.append({
                        'type': 'CMS Detection',
                        'severity': 'Low',
                        'cvss_score': 3.1,
                        'cvss_vector': 'CVSS:3.1/AV:N/AC:H/PR:N/UI:R/S:U/C:L/I:N/A:N',
                        'category': 'A06',
                        'owasp_category': 'A06:2021 - Vulnerable and Outdated Components',
                        'url': self.target_url,
                        'cms': cms_name,
                        'description': f'{cms_name} CMS detected. Ensure it is up-to-date and patched.',
                        'impact': 'Potential exploitation if CMS contains known vulnerabilities',
                        'remediation': f'Keep {cms_name} updated to latest version and apply security patches'
                    })
        
        except:
            pass
        
        return vulnerabilities

    def scan_a07_identification_auth_failures(self):
        """A07:2021 – Identification and Authentication Failures"""
        vulnerabilities = []
        
        try:
            # Test for common weak credentials
            login_paths = ['/login', '/admin', '/signin', '/auth']
            weak_creds = [
                ('admin', 'admin'),
                ('admin', 'password'),
                ('root', 'root'),
                ('test', 'test'),
                ('guest', 'guest')
            ]
            
            for path in login_paths:
                login_url = self.target_url + path
                try:
                    response = self.session.get(login_url, timeout=5)
                    if response.status_code == 200 and 'login' in response.text.lower():
                        
                        # Test weak credentials
                        for username, password in weak_creds:
                            try:
                                login_data = {'username': username, 'password': password}
                                auth_response = self.session.post(login_url, data=login_data, timeout=5)
                                
                                if 'welcome' in auth_response.text.lower() or 'dashboard' in auth_response.text.lower():
                                    vulnerabilities.append({
                                        'type': 'Weak Default Credentials',
                                        'severity': 'Critical',
                                        'cvss_score': 9.8,
                                        'cvss_vector': 'CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H',
                                        'category': 'A07',
                                        'owasp_category': 'A07:2021 - Identification and Authentication Failures',
                                        'url': login_url,
                                        'username': username,
                                        'description': f'System accepts weak credentials: {username}/{password}',
                                        'impact': 'Complete system compromise, unauthorized access to all functions',
                                        'remediation': 'Remove default accounts, enforce strong password policies'
                                    })
                                
                                time.sleep(self.delay)
                            except:
                                continue
                        
                        # Check for account enumeration
                        enum_response = self.session.post(login_url, data={'username': 'nonexistentuser123456', 'password': 'wrongpass'}, timeout=5)
                        if 'user not found' in enum_response.text.lower() or 'invalid username' in enum_response.text.lower():
                            vulnerabilities.append({
                                'type': 'Username Enumeration',
                                'severity': 'Medium',
                                'cvss_score': 5.3,
                                'cvss_vector': 'CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:L/I:N/A:N',
                                'category': 'A07',
                                'owasp_category': 'A07:2021 - Identification and Authentication Failures',
                                'url': login_url,
                                'description': 'Login form reveals whether usernames exist, enabling user enumeration attacks',
                                'impact': 'Account enumeration, targeted brute force attacks',
                                'remediation': 'Use generic error messages that do not reveal username validity'
                            })
                
                except:
                    continue
        
        except:
            pass
        
        return vulnerabilities

    def scan_a08_software_data_integrity_failures(self):
        """A08:2021 – Software and Data Integrity Failures"""
        vulnerabilities = []
        
        try:
            response = self.session.get(self.target_url, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Check for external resources without integrity checks
            external_scripts = soup.find_all('script', src=True)
            external_links = soup.find_all('link', href=True)
            
            for script in external_scripts:
                src = script.get('src', '')
                if ('://' in src and self.target_url not in src) and not script.get('integrity'):
                        vulnerabilities.append({
                            'type': 'Missing Subresource Integrity',
                            'severity': 'Medium',
                            'cvss_score': 6.1,
                            'cvss_vector': 'CVSS:3.1/AV:N/AC:L/PR:N/UI:R/S:C/C:L/I:L/A:N',
                            'category': 'A08',
                            'owasp_category': 'A08:2021 - Software and Data Integrity Failures',
                            'url': self.target_url,
                            'resource': src,
                            'description': f'External script loaded without integrity check: {src}',
                            'impact': 'Potential code injection if external resource is compromised',
                            'remediation': 'Add integrity attribute with cryptographic hash to external resources'
                        })
            
            for link in external_links:
                href = link.get('href', '')
                if ('://' in href and self.target_url not in href) and not link.get('integrity'):
                    vulnerabilities.append({
                        'type': 'Missing Subresource Integrity',
                        'severity': 'Low',
                        'cvss_score': 4.3,
                        'cvss_vector': 'CVSS:3.1/AV:N/AC:L/PR:N/UI:R/S:U/C:L/I:N/A:N',
                        'category': 'A08',
                        'owasp_category': 'A08:2021 - Software and Data Integrity Failures',
                        'url': self.target_url,
                        'resource': href,
                        'description': f'External resource loaded without integrity check: {href}',
                        'impact': 'Potential resource tampering if external source is compromised',
                        'remediation': 'Add integrity attribute with cryptographic hash to external stylesheets'
                    })
        
        except:
            pass
        
        return vulnerabilities

    def scan_a09_security_logging_failures(self):
        """A09:2021 – Security Logging and Monitoring Failures"""
        vulnerabilities = []
        
        # Check for missing security logging
        try:
            # Test failed login attempt
            response = self.session.post(f"{self.target_url}/login", 
                                       data={'username': 'invalid_user_test', 'password': 'invalid_pass_test'}, timeout=5)
            
            # Check for missing logging indicators
            logging_headers = ['x-request-id', 'x-correlation-id', 'x-trace-id', 'x-log-id']
            has_logging_headers = any(header in response.headers for header in logging_headers)
            
            if not has_logging_headers:
                vulnerabilities.append({
                    'type': 'Missing Security Event Logging',
                    'severity': 'Medium',
                    'cvss_score': 6.1,
                    'cvss_vector': 'CVSS:3.1/AV:N/AC:L/PR:N/UI:R/S:C/C:L/I:L/A:N',
                    'category': 'A09',
                    'owasp_category': 'A09:2021 - Security Logging and Monitoring Failures',
                    'url': f"{self.target_url}/login",
                    'description': 'No evidence of security event logging detected. Missing request tracking headers and audit trail.',
                    'impact': 'Delayed incident response, difficulty in forensic analysis, compliance violations',
                    'remediation': 'Implement comprehensive security logging, audit trails, and monitoring systems'
                })
            
            # Check for verbose error messages that might indicate lack of proper logging
            error_indicators = ['stack trace', 'exception', 'error occurred', 'debug', 'traceback', 'mysql_connect', 'postgresql', 'ora-', 'sqlite_']
            if any(indicator in response.text.lower() for indicator in error_indicators):
                vulnerabilities.append({
                    'type': 'Verbose Error Messages',
                    'severity': 'Low',
                    'cvss_score': 4.3,
                    'cvss_vector': 'CVSS:3.1/AV:N/AC:L/PR:N/UI:R/S:U/C:L/I:N/A:N',
                    'category': 'A09',
                    'owasp_category': 'A09:2021 - Security Logging and Monitoring Failures',
                    'url': self.target_url,
                    'description': 'Application exposes detailed error messages that could aid attackers',
                    'impact': 'Information disclosure, system reconnaissance, attack surface mapping, database technology disclosure',
                    'remediation': 'Implement proper error handling with generic user messages and detailed server-side logging'
                })
                
            # Check for missing rate limiting
            try:
                # Test multiple requests quickly
                for i in range(5):
                    quick_response = self.session.get(self.target_url, timeout=3)
                    if quick_response.status_code != 429:  # No rate limiting
                        pass
                else:
                    # If we get here without 429, likely no rate limiting
                    vulnerabilities.append({
                        'type': 'Missing Rate Limiting',
                        'severity': 'Medium',
                        'cvss_score': 5.3,
                        'cvss_vector': 'CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:N/I:N/A:L',
                        'category': 'A09',
                        'owasp_category': 'A09:2021 - Security Logging and Monitoring Failures',
                        'url': self.target_url,
                        'description': 'No rate limiting detected on repeated requests',
                        'impact': 'Vulnerable to brute force attacks, DoS attacks, resource exhaustion',
                        'remediation': 'Implement rate limiting and request throttling mechanisms'
                    })
            except:
                pass
        
        except Exception as e:
            pass
            
        # Check for exposed admin panels
        admin_paths = [
            '/admin', '/administrator', '/admin.php', '/admin/', '/wp-admin',
            '/phpmyadmin', '/cpanel', '/control', '/manager', '/admin/login',
            '/backend', '/dashboard', '/console', '/panel'
        ]
        
        for path in admin_paths:
            try:
                admin_url = f"{self.target_url}{path}"
                response = self.session.get(admin_url, timeout=5)
                if response.status_code == 200 and 'login' in response.text.lower():
                    vulnerabilities.append({
                        'type': 'Exposed Admin Panel',
                        'severity': 'High',
                        'cvss_score': 7.5,
                        'cvss_vector': 'CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N',
                        'category': 'A01',
                        'owasp_category': 'A01:2021 - Broken Access Control',
                        'url': admin_url,
                        'description': f'Admin panel accessible at {path}',
                        'impact': 'Unauthorized administrative access if credentials are weak',
                        'remediation': 'Restrict admin panel access by IP, implement strong authentication'
                    })
            except:
                continue
        
        return vulnerabilities

    def scan_a10_ssrf(self):
        """A10:2021 – Server-Side Request Forgery"""
        vulnerabilities = []
        
        # SSRF test payloads
        ssrf_payloads = [
            'http://127.0.0.1/',
            'http://localhost/',
            'http://169.254.169.254/',  # AWS metadata
            'http://metadata.google.internal/',  # GCP metadata
            'file:///etc/passwd',
            'http://0.0.0.0/',
            'http://[::1]/'
        ]
        
        try:
            response = self.session.get(self.target_url, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Find forms that might accept URLs
            forms = soup.find_all('form')
            for form in forms:
                inputs = form.find_all('input')
                
                for input_field in inputs:
                    field_name = input_field.get('name', '')
                    field_type = input_field.get('type', 'text')
                    
                    # Look for URL-like field names
                    if any(keyword in field_name.lower() for keyword in ['url', 'link', 'callback', 'redirect', 'fetch']):
                        form_url = form.get('action', self.target_url)
                        if not form_url.startswith('http'):
                            form_url = self.target_url + form_url if form_url.startswith('/') else self.target_url + '/' + form_url
                        
                        # Test SSRF payloads
                        for payload in ssrf_payloads[:3]:  # Test first 3 payloads
                            try:
                                test_data = {field_name: payload}
                                
                                if form.get('method', '').lower() == 'post':
                                    ssrf_response = self.session.post(form_url, data=test_data, timeout=5)
                                else:
                                    ssrf_response = self.session.get(form_url, params=test_data, timeout=5)
                                
                                # Check for SSRF indicators
                                if any(indicator in ssrf_response.text.lower() for indicator in ['root:', 'ubuntu', 'centos', 'localhost']):
                                    vulnerabilities.append({
                                        'type': 'Potential SSRF',
                                        'severity': 'High',
                                        'cvss_score': 8.6,
                                        'cvss_vector': 'CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:C/C:H/I:N/A:N',
                                        'category': 'A10',
                                        'owasp_category': 'A10:2021 - Server-Side Request Forgery',
                                        'url': form_url,
                                        'parameter': field_name,
                                        'payload': payload,
                                        'description': f'Possible SSRF vulnerability in parameter {field_name}',
                                        'impact': 'Access to internal networks, cloud metadata exposure, port scanning',
                                        'remediation': 'Validate and whitelist allowed URLs, implement network segmentation'
                                    })
                                
                                time.sleep(self.delay)
                            except:
                                continue
        
        except:
            pass
        
        return vulnerabilities

    def comprehensive_scan(self):
        """Run comprehensive OWASP Top 10 scan"""
        self.crawl_website()
        
        owasp_scans = [
            ('A01:2021 - Broken Access Control', self.scan_a01_broken_access_control),
            ('A02:2021 - Cryptographic Failures', self.scan_a02_cryptographic_failures),
            ('A03:2021 - Injection', self.scan_a03_injection),
            ('A04:2021 - Insecure Design', self.scan_a04_insecure_design),
            ('A05:2021 - Security Misconfiguration', self.scan_a05_security_misconfiguration),
            ('A06:2021 - Vulnerable and Outdated Components', self.scan_a06_vulnerable_components),
            ('A07:2021 - Identification and Authentication Failures', self.scan_a07_identification_auth_failures),
            ('A08:2021 - Software and Data Integrity Failures', self.scan_a08_software_data_integrity_failures),
            ('A09:2021 - Security Logging and Monitoring Failures', self.scan_a09_security_logging_failures),
            ('A10:2021 - Server-Side Request Forgery', self.scan_a10_ssrf)
        ]
        
        all_vulnerabilities = []
        
        for category_name, scan_function in owasp_scans:
            try:
                vulnerabilities = scan_function()
                all_vulnerabilities.extend(vulnerabilities)
            except Exception as e:
                pass  # Continue with other scans even if one fails
        
        self.vulnerabilities = all_vulnerabilities
        return all_vulnerabilities
    
    def export_enhanced_results(self):
        """Export comprehensive scan results with detailed reporting"""
        return {
            'scan_metadata': {
                'target_url': self.target_url,
                'scan_date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'scanner_version': VERSION,
                'total_vulnerabilities': len(self.vulnerabilities),
                'scan_parameters': {
                    'max_depth': self.max_depth,
                    'delay': self.delay,
                    'urls_scanned': len(self.visited_urls)
                }
            },
            'executive_summary': self.generate_executive_summary(),
            'vulnerability_breakdown': self.generate_vulnerability_breakdown(),
            'detailed_findings': self.vulnerabilities,
            'owasp_coverage': self.get_owasp_coverage_report(),
            'recommendations': self.generate_recommendations(),
            'risk_matrix': self.generate_risk_matrix()
        }
    
    def generate_executive_summary(self):
        """Generate executive summary for the scan results"""
        total_vulns = len(self.vulnerabilities)
        if total_vulns == 0:
            return {
                'status': 'SECURE',
                'risk_level': 'LOW',
                'summary': f'No critical vulnerabilities detected on {self.target_url}. The application appears to follow basic security practices.',
                'key_findings': ['No OWASP Top 10 vulnerabilities detected', 'Application follows basic security guidelines']
            }
        
        # Count by severity
        severity_counts = {'Critical': 0, 'High': 0, 'Medium': 0, 'Low': 0}
        for vuln in self.vulnerabilities:
            severity = vuln.get('severity', 'Low')
            severity_counts[severity] = severity_counts.get(severity, 0) + 1
        
        # Determine overall risk level
        if severity_counts['Critical'] > 0:
            risk_level = 'CRITICAL'
            status = 'VULNERABLE'
        elif severity_counts['High'] > 3:
            risk_level = 'HIGH'
            status = 'VULNERABLE'
        elif severity_counts['High'] > 0 or severity_counts['Medium'] > 5:
            risk_level = 'MEDIUM'
            status = 'AT_RISK'
        else:
            risk_level = 'LOW'
            status = 'MINOR_ISSUES'
        
        # Generate key findings
        key_findings = []
        if severity_counts['Critical'] > 0:
            key_findings.append(f'{severity_counts["Critical"]} Critical vulnerabilities require immediate attention')
        if severity_counts['High'] > 0:
            key_findings.append(f'{severity_counts["High"]} High-severity issues need prompt remediation')
        
        # Identify most common vulnerability types
        vuln_types = {}
        for vuln in self.vulnerabilities:
            vuln_type = vuln.get('type', 'Unknown')
            vuln_types[vuln_type] = vuln_types.get(vuln_type, 0) + 1
        
        if vuln_types:
            most_common = max(vuln_types.items(), key=lambda x: x[1])
            key_findings.append(f'Most common issue: {most_common[0]} ({most_common[1]} instances)')
        
        return {
            'status': status,
            'risk_level': risk_level,
            'summary': f'Security assessment of {self.target_url} identified {total_vulns} vulnerabilities across {len(vuln_types)} different categories.',
            'key_findings': key_findings,
            'severity_distribution': severity_counts
        }
    
    def generate_vulnerability_breakdown(self):
        """Generate detailed vulnerability breakdown by category and type"""
        breakdown = {
            'by_owasp_category': {},
            'by_vulnerability_type': {},
            'by_severity': {'Critical': 0, 'High': 0, 'Medium': 0, 'Low': 0},
            'by_cvss_score': {'9.0-10.0': 0, '7.0-8.9': 0, '4.0-6.9': 0, '0.1-3.9': 0}
        }
        
        for vuln in self.vulnerabilities:
            # Count by OWASP category
            category = vuln.get('category', 'Unknown')
            if category not in breakdown['by_owasp_category']:
                breakdown['by_owasp_category'][category] = {'count': 0, 'vulnerabilities': []}
            breakdown['by_owasp_category'][category]['count'] += 1
            breakdown['by_owasp_category'][category]['vulnerabilities'].append({
                'type': vuln.get('type'),
                'severity': vuln.get('severity'),
                'cvss_score': vuln.get('cvss_score')
            })
            
            # Count by vulnerability type
            vuln_type = vuln.get('type', 'Unknown')
            breakdown['by_vulnerability_type'][vuln_type] = breakdown['by_vulnerability_type'].get(vuln_type, 0) + 1
            
            # Count by severity
            severity = vuln.get('severity', 'Low')
            breakdown['by_severity'][severity] += 1
            
            # Count by CVSS score range
            cvss_score = vuln.get('cvss_score', 0)
            if cvss_score >= 9.0:
                breakdown['by_cvss_score']['9.0-10.0'] += 1
            elif cvss_score >= 7.0:
                breakdown['by_cvss_score']['7.0-8.9'] += 1
            elif cvss_score >= 4.0:
                breakdown['by_cvss_score']['4.0-6.9'] += 1
            else:
                breakdown['by_cvss_score']['0.1-3.9'] += 1
        
        return breakdown
    
    def get_owasp_coverage_report(self):
        """Generate OWASP Top 10 coverage report with findings"""
        owasp_categories = self.get_owasp_top10_coverage()
        coverage_report = {}
        
        for category_id, category_info in owasp_categories.items():
            # Count vulnerabilities found in this category
            category_vulns = [v for v in self.vulnerabilities if v.get('category') == category_id]
            
            coverage_report[category_id] = {
                'name': category_info['name'],
                'description': category_info['description'],
                'vulnerabilities_found': len(category_vulns),
                'tested': True,  # All categories are tested
                'status': 'VULNERABLE' if category_vulns else 'SECURE',
                'findings': [{
                    'type': v.get('type'),
                    'severity': v.get('severity'),
                    'cvss_score': v.get('cvss_score')
                } for v in category_vulns]
            }
        
        return coverage_report
    
    def generate_recommendations(self):
        """Generate prioritized recommendations based on findings"""
        recommendations = []
        
        # Count vulnerabilities by type and severity
        critical_vulns = [v for v in self.vulnerabilities if v.get('severity') == 'Critical']
        high_vulns = [v for v in self.vulnerabilities if v.get('severity') == 'High']
        
        if critical_vulns:
            recommendations.append({
                'priority': 'IMMEDIATE',
                'title': 'Address Critical Vulnerabilities',
                'description': f'Immediately remediate {len(critical_vulns)} critical vulnerabilities that pose severe security risks.',
                'affected_areas': list(set([v.get('type') for v in critical_vulns])),
                'timeline': '0-24 hours'
            })
        
        if high_vulns:
            recommendations.append({
                'priority': 'HIGH',
                'title': 'Remediate High-Risk Issues',
                'description': f'Address {len(high_vulns)} high-severity vulnerabilities within the next week.',
                'affected_areas': list(set([v.get('type') for v in high_vulns])),
                'timeline': '1-7 days'
            })
        
        # General security recommendations
        recommendations.extend([
            {
                'priority': 'MEDIUM',
                'title': 'Implement Security Headers',
                'description': 'Add comprehensive security headers to prevent common attacks.',
                'affected_areas': ['HTTP Security', 'XSS Prevention', 'Clickjacking'],
                'timeline': '1-2 weeks'
            },
            {
                'priority': 'MEDIUM',
                'title': 'Regular Security Updates',
                'description': 'Establish a process for regular security updates and vulnerability management.',
                'affected_areas': ['Dependency Management', 'Patch Management'],
                'timeline': 'Ongoing'
            },
            {
                'priority': 'LOW',
                'title': 'Security Monitoring',
                'description': 'Implement comprehensive security logging and monitoring.',
                'affected_areas': ['Logging', 'Monitoring', 'Incident Response'],
                'timeline': '1 month'
            }
        ])
        
        return recommendations
    
    def generate_risk_matrix(self):
        """Generate risk assessment matrix"""
        risk_matrix = {
            'overall_risk_score': 0,
            'risk_factors': [],
            'mitigation_priority': []
        }
        
        # Calculate overall risk score based on CVSS scores
        total_risk = sum([v.get('cvss_score', 0) for v in self.vulnerabilities])
        risk_matrix['overall_risk_score'] = round(total_risk / max(len(self.vulnerabilities), 1), 2)
        
        # Identify key risk factors
        vuln_types = {}
        for vuln in self.vulnerabilities:
            vuln_type = vuln.get('type', 'Unknown')
            if vuln_type not in vuln_types:
                vuln_types[vuln_type] = {'count': 0, 'max_cvss': 0}
            vuln_types[vuln_type]['count'] += 1
            vuln_types[vuln_type]['max_cvss'] = max(vuln_types[vuln_type]['max_cvss'], vuln.get('cvss_score', 0))
        
        # Sort by risk (count * max CVSS)
        sorted_risks = sorted(vuln_types.items(), key=lambda x: x[1]['count'] * x[1]['max_cvss'], reverse=True)
        
        for vuln_type, data in sorted_risks[:5]:  # Top 5 risks
            risk_matrix['risk_factors'].append({
                'vulnerability_type': vuln_type,
                'count': data['count'],
                'max_cvss_score': data['max_cvss'],
                'risk_level': 'Critical' if data['max_cvss'] >= 9.0 else 'High' if data['max_cvss'] >= 7.0 else 'Medium'
            })
        
        # Generate mitigation priorities
        critical_items = [v for v in self.vulnerabilities if v.get('cvss_score', 0) >= 9.0]
        high_items = [v for v in self.vulnerabilities if 7.0 <= v.get('cvss_score', 0) < 9.0]
        
        if critical_items:
            risk_matrix['mitigation_priority'].append('Immediately address critical vulnerabilities')
        if high_items:
            risk_matrix['mitigation_priority'].append('Prioritize high-severity issues')
        if len(self.vulnerabilities) > 10:
            risk_matrix['mitigation_priority'].append('Develop systematic remediation plan')
        
        return risk_matrix
    
    def generate_html_report(self):
        """Generate a comprehensive HTML report"""
        enhanced_data = self.export_enhanced_results()
        
        html_report = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Security Assessment Report - {self.target_url}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            margin: 0;
            padding: 20px;
            background: #f8f9fa;
            color: #333;
        }}
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            padding: 30px;
            border-radius: 12px;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }}
        .header {{
            text-align: center;
            margin-bottom: 40px;
            border-bottom: 3px solid #667eea;
            padding-bottom: 20px;
        }}
        .header h1 {{
            color: #667eea;
            font-size: 28px;
            margin: 0;
        }}
        .header .subtitle {{
            color: #666;
            font-size: 18px;
            margin: 10px 0;
        }}
        .section {{
            margin-bottom: 30px;
        }}
        .section h2 {{
            color: #667eea;
            border-left: 4px solid #667eea;
            padding-left: 15px;
            margin-bottom: 15px;
        }}
        .executive-summary {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 25px;
            border-radius: 8px;
            margin-bottom: 30px;
        }}
        .risk-badge {{
            display: inline-block;
            padding: 8px 16px;
            border-radius: 20px;
            font-weight: bold;
            font-size: 14px;
        }}
        .risk-critical {{ background: #dc3545; color: white; }}
        .risk-high {{ background: #fd7e14; color: white; }}
        .risk-medium {{ background: #ffc107; color: black; }}
        .risk-low {{ background: #28a745; color: white; }}
        .vuln-item {{
            background: #f8f9fa;
            border-left: 4px solid #dc3545;
            padding: 15px;
            margin-bottom: 15px;
            border-radius: 4px;
        }}
        .vuln-item.high {{ border-left-color: #fd7e14; }}
        .vuln-item.medium {{ border-left-color: #ffc107; }}
        .vuln-item.low {{ border-left-color: #28a745; }}
        .stats-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
            margin: 20px 0;
        }}
        .stat-card {{
            background: #f8f9fa;
            padding: 20px;
            border-radius: 8px;
            text-align: center;
        }}
        .stat-number {{
            font-size: 32px;
            font-weight: bold;
            color: #667eea;
        }}
        .recommendations {{
            background: #e8f2ff;
            padding: 20px;
            border-radius: 8px;
            border-left: 4px solid #667eea;
        }}
        .recommendation-item {{
            margin-bottom: 15px;
            padding: 15px;
            background: white;
            border-radius: 6px;
        }}
        .priority-immediate {{ border-left: 4px solid #dc3545; }}
        .priority-high {{ border-left: 4px solid #fd7e14; }}
        .priority-medium {{ border-left: 4px solid #ffc107; }}
        .priority-low {{ border-left: 4px solid #28a745; }}
        @media print {{
            body {{ background: white; }}
            .container {{ box-shadow: none; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>🛡️ Security Assessment Report</h1>
            <div class="subtitle">Target: {self.target_url}</div>
            <div class="subtitle">Generated: {enhanced_data['scan_metadata']['scan_date']}</div>
        </div>
        
        <div class="executive-summary">
            <h2 style="color: white; border: none; padding: 0;">📋 Executive Summary</h2>
            <div style="display: flex; align-items: center; gap: 20px; flex-wrap: wrap;">
                <div>
                    <div class="risk-badge risk-{enhanced_data['executive_summary']['risk_level'].lower()}">
                        {enhanced_data['executive_summary']['risk_level']} RISK
                    </div>
                </div>
                <div style="flex: 1;">                    
                    <p>{enhanced_data['executive_summary']['summary']}</p>
                </div>
            </div>
        </div>
        
        <div class="stats-grid">
            <div class="stat-card">
                <div class="stat-number">{enhanced_data['scan_metadata']['total_vulnerabilities']}</div>
                <div>Total Vulnerabilities</div>
            </div>
            <div class="stat-card">
                <div class="stat-number">{enhanced_data['executive_summary']['severity_distribution']['Critical']}</div>
                <div>Critical Issues</div>
            </div>
            <div class="stat-card">
                <div class="stat-number">{enhanced_data['executive_summary']['severity_distribution']['High']}</div>
                <div>High Risk Issues</div>
            </div>
            <div class="stat-card">
                <div class="stat-number">{len(enhanced_data['owasp_coverage'])}</div>
                <div>OWASP Categories Tested</div>
            </div>
        </div>"""
        
        # Add vulnerability details
        if self.vulnerabilities:
            html_report += """
        <div class="section">
            <h2>🔍 Vulnerability Details</h2>"""
            
            for i, vuln in enumerate(self.vulnerabilities[:10]):  # Limit to top 10 for readability
                severity_class = vuln.get('severity', 'low').lower()
                html_report += f"""
            <div class="vuln-item {severity_class}">
                <h3>{vuln.get('type', 'Unknown Vulnerability')}</h3>
                <p><strong>Severity:</strong> <span class="risk-badge risk-{severity_class}">{vuln.get('severity', 'Unknown')}</span></p>
                <p><strong>CVSS Score:</strong> {vuln.get('cvss_score', 'N/A')}</p>
                <p><strong>Description:</strong> {vuln.get('description', 'No description available')}</p>
                {f'<p><strong>URL:</strong> {vuln.get("url", "N/A")}</p>' if vuln.get('url') else ''}
                {f'<p><strong>Remediation:</strong> {vuln.get("remediation", "Consult security documentation")}</p>' if vuln.get('remediation') else ''}
            </div>"""
            
            html_report += "</div>"
        
        # Add recommendations
        html_report += f"""
        <div class="section">
            <h2>💡 Recommendations</h2>
            <div class="recommendations">"""
        
        for rec in enhanced_data['recommendations']:
            priority_class = rec['priority'].lower().replace('_', '-')
            html_report += f"""
                <div class="recommendation-item priority-{priority_class}">
                    <h4>{rec['title']} ({rec['priority']} Priority)</h4>
                    <p>{rec['description']}</p>
                    <p><strong>Timeline:</strong> {rec['timeline']}</p>
                </div>"""
        
        html_report += """
            </div>
        </div>
        
        <div class="section">
            <h2>📊 OWASP Top 10 Coverage</h2>
            <div class="stats-grid">"""
        
        for category_id, category_data in enhanced_data['owasp_coverage'].items():
            status_class = 'success' if category_data['status'] == 'SECURE' else 'danger'
            html_report += f"""
                <div class="stat-card">
                    <h4>{category_id}</h4>
                    <p>{category_data['name']}</p>
                    <div class="risk-badge risk-{'low' if category_data['status'] == 'SECURE' else 'high'}">
                        {category_data['vulnerabilities_found']} Issues
                    </div>
                </div>"""
        
        html_report += """
            </div>
        </div>
        
        <div style="margin-top: 40px; padding-top: 20px; border-top: 1px solid #ddd; text-align: center; color: #666;">
            <p>Report generated by OWASP Top 10 Vulnerability Scanner v2.1.0</p>
            <p><strong>Disclaimer:</strong> This report is for authorized security testing only. 
            Use findings responsibly and in accordance with applicable laws and regulations.</p>
        </div>
    </div>
</body>
</html>"""
        
        return html_report
    
    def generate_json_report(self):
        """Generate machine-readable JSON report"""
        return self.export_enhanced_results()
    
    def generate_pdf_report(self):
        """Generate PDF report using reportlab for cross-platform compatibility"""
        try:
            # Use reportlab for reliable PDF generation on all platforms
            from reportlab.lib.pagesizes import A4, letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
            from io import BytesIO
            
            # Create PDF buffer
            buffer = BytesIO()
            
            # Create PDF document
            doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # Get enhanced data
            enhanced_data = self.export_enhanced_results()
            
            # Build story (content)
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=TA_CENTER,
                textColor=colors.HexColor('#2E5BBA')
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                spaceAfter=12,
                textColor=colors.HexColor('#2E5BBA')
            )
            
            subheading_style = ParagraphStyle(
                'CustomSubheading',
                parent=styles['Heading3'],
                fontSize=14,
                spaceAfter=8,
                textColor=colors.HexColor('#666666')
            )
            
            # Title page
            story.append(Paragraph('🛡️ Security Assessment Report', title_style))
            story.append(Spacer(1, 20))
            story.append(Paragraph(f'Target: {self.target_url}', styles['Normal']))
            story.append(Paragraph(f'Generated: {enhanced_data["scan_metadata"]["scan_date"]}', styles['Normal']))
            story.append(Paragraph(f'Scanner Version: {enhanced_data["scan_metadata"]["scanner_version"]}', styles['Normal']))
            story.append(PageBreak())
            
            # Executive Summary
            story.append(Paragraph('Executive Summary', heading_style))
            exec_summary = enhanced_data['executive_summary']
            
            # Risk level with color coding
            risk_level = exec_summary['risk_level']
            risk_color = colors.red if risk_level == 'CRITICAL' else colors.orange if risk_level == 'HIGH' else colors.yellow if risk_level == 'MEDIUM' else colors.green
            
            story.append(Paragraph(f'<b>Risk Level: <font color="{risk_color}">{risk_level}</font></b>', styles['Normal']))
            story.append(Spacer(1, 12))
            story.append(Paragraph(exec_summary['summary'], styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Key findings
            story.append(Paragraph('Key Findings:', subheading_style))
            for finding in exec_summary['key_findings']:
                story.append(Paragraph(f'• {finding}', styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Vulnerability Statistics Table
            story.append(Paragraph('Vulnerability Statistics', heading_style))
            
            severity_data = [['Severity Level', 'Count']]
            for severity, count in exec_summary['severity_distribution'].items():
                severity_data.append([severity, str(count)])
            
            severity_table = Table(severity_data)
            severity_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E5BBA')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(severity_table)
            story.append(Spacer(1, 20))
            
            # Detailed Findings
            if self.vulnerabilities:
                story.append(Paragraph('Detailed Vulnerability Findings', heading_style))
                
                for i, vuln in enumerate(self.vulnerabilities[:15], 1):  # Limit for PDF size
                    story.append(Paragraph(f'{i}. {vuln.get("type", "Unknown Vulnerability")}', subheading_style))
                    
                    # Vulnerability details table
                    vuln_data = [
                        ['Severity', vuln.get('severity', 'Unknown')],
                        ['CVSS Score', str(vuln.get('cvss_score', 'N/A'))],
                    ]
                    
                    if vuln.get('url'):
                        vuln_data.append(['Location', vuln.get('url')])
                    
                    vuln_table = Table(vuln_data, colWidths=[2*inch, 4*inch])
                    vuln_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F0F0F0')),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 10),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(vuln_table)
                    
                    if vuln.get('description'):
                        story.append(Paragraph(f'<b>Description:</b> {vuln.get("description")}', styles['Normal']))
                    
                    if vuln.get('remediation'):
                        story.append(Paragraph(f'<b>Remediation:</b> {vuln.get("remediation")}', styles['Normal']))
                    
                    story.append(Spacer(1, 15))
            
            # Recommendations
            story.append(PageBreak())
            story.append(Paragraph('Security Recommendations', heading_style))
            
            for rec in enhanced_data['recommendations']:
                priority_color = colors.red if rec['priority'] == 'IMMEDIATE' else colors.orange if rec['priority'] == 'HIGH' else colors.blue
                story.append(Paragraph(f'<font color="{priority_color}"><b>{rec["title"]} ({rec["priority"]} Priority)</b></font>', subheading_style))
                story.append(Paragraph(rec['description'], styles['Normal']))
                story.append(Paragraph(f'<b>Timeline:</b> {rec["timeline"]}', styles['Normal']))
                story.append(Spacer(1, 10))
            
            # OWASP Coverage
            story.append(PageBreak())
            story.append(Paragraph('OWASP Top 10 Coverage Analysis', heading_style))
            
            owasp_data = [['Category', 'Name', 'Vulnerabilities', 'Status']]
            for category_id, category_data in enhanced_data['owasp_coverage'].items():
                status = 'SECURE' if category_data['vulnerabilities_found'] == 0 else 'VULNERABLE'
                owasp_data.append([
                    category_id,
                    category_data['name'][:30] + '...' if len(category_data['name']) > 30 else category_data['name'],
                    str(category_data['vulnerabilities_found']),
                    status
                ])
            
            owasp_table = Table(owasp_data, colWidths=[1*inch, 2.5*inch, 1*inch, 1*inch])
            owasp_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E5BBA')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(owasp_table)
            
            # Footer
            story.append(Spacer(1, 30))
            story.append(Paragraph('Report generated by OWASP Top 10 Vulnerability Scanner v2.1.0', styles['Normal']))
            story.append(Paragraph('<b>Disclaimer:</b> This report is for authorized security testing only. Use findings responsibly and in accordance with applicable laws.', styles['Normal']))
            
            # Build PDF
            doc.build(story)
            
            # Get PDF data
            pdf_data = buffer.getvalue()
            buffer.close()
            
            return pdf_data
            
        except ImportError:
            # Fallback: Generate structured text report as PDF-ready content
            enhanced_data = self.export_enhanced_results()
            
            text_content = f"""SECURITY ASSESSMENT REPORT
{'='*60}

Target: {self.target_url}
Generated: {enhanced_data['scan_metadata']['scan_date']}
Scanner Version: {enhanced_data['scan_metadata']['scanner_version']}

EXECUTIVE SUMMARY
{'-'*20}
Risk Level: {enhanced_data['executive_summary']['risk_level']}
Status: {enhanced_data['executive_summary']['status']}

{enhanced_data['executive_summary']['summary']}

KEY FINDINGS:
"""
            
            for finding in enhanced_data['executive_summary']['key_findings']:
                text_content += f"• {finding}\n"
            
            text_content += f"\n\nVULNERABILITY BREAKDOWN\n{'-'*25}\n"
            
            for severity, count in enhanced_data['executive_summary']['severity_distribution'].items():
                text_content += f"{severity}: {count}\n"
            
            if self.vulnerabilities:
                text_content += f"\n\nDETAILED FINDINGS\n{'-'*20}\n"
                
                for i, vuln in enumerate(self.vulnerabilities[:10], 1):
                    text_content += f"\n{i}. {vuln.get('type', 'Unknown')}\n"
                    text_content += f"   Severity: {vuln.get('severity', 'Unknown')}\n"
                    text_content += f"   CVSS Score: {vuln.get('cvss_score', 'N/A')}\n"
                    if vuln.get('description'):
                        text_content += f"   Description: {vuln.get('description')}\n"
                    if vuln.get('remediation'):
                        text_content += f"   Remediation: {vuln.get('remediation')}\n"
            
            text_content += f"\n\nRECOMMENDATIONS\n{'-'*15}\n"
            
            for rec in enhanced_data['recommendations']:
                text_content += f"\n{rec['priority']} PRIORITY: {rec['title']}\n"
                text_content += f"Timeline: {rec['timeline']}\n"
                text_content += f"{rec['description']}\n"
            
            return text_content.encode('utf-8')
        
        except Exception as e:
            # Final fallback: Simple text report
            enhanced_data = self.export_enhanced_results()
            fallback_content = f"""Security Assessment Report\n\nTarget: {self.target_url}\nDate: {enhanced_data['scan_metadata']['scan_date']}\n\nRisk Level: {enhanced_data['executive_summary']['risk_level']}\nTotal Vulnerabilities: {enhanced_data['scan_metadata']['total_vulnerabilities']}\n\nThis is a fallback text report. Please install 'reportlab' for full PDF generation.\n\nError: {str(e)}"""
            return fallback_content.encode('utf-8')
    
    def generate_word_report(self):
        """Generate Word document report"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
            
            # Create new document
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Get enhanced data
            enhanced_data = self.export_enhanced_results()
            
            # Title page
            title = doc.add_heading('Security Assessment Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Subtitle
            subtitle = doc.add_paragraph(f'Target: {self.target_url}')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            date_para = doc.add_paragraph(f'Generated: {enhanced_data["scan_metadata"]["scan_date"]}')
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_page_break()
            
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            
            exec_summary = enhanced_data['executive_summary']
            doc.add_paragraph(f'Risk Level: {exec_summary["risk_level"]}')
            doc.add_paragraph(exec_summary['summary'])
            
            # Key findings
            doc.add_heading('Key Findings', level=2)
            for finding in exec_summary['key_findings']:
                p = doc.add_paragraph(finding, style='List Bullet')
            
            # Vulnerability Statistics
            doc.add_heading('Vulnerability Statistics', level=2)
            
            # Create table for severity distribution
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Severity Level'
            hdr_cells[1].text = 'Count'
            
            for severity, count in exec_summary['severity_distribution'].items():
                row_cells = table.add_row().cells
                row_cells[0].text = severity
                row_cells[1].text = str(count)
            
            # Detailed Findings
            if self.vulnerabilities:
                doc.add_heading('Detailed Vulnerability Findings', level=1)
                
                for i, vuln in enumerate(self.vulnerabilities[:20], 1):  # Limit to top 20
                    doc.add_heading(f'{i}. {vuln.get("type", "Unknown Vulnerability")}', level=2)
                    
                    # Vulnerability details
                    doc.add_paragraph(f'Severity: {vuln.get("severity", "Unknown")}')
                    doc.add_paragraph(f'CVSS Score: {vuln.get("cvss_score", "N/A")}')
                    
                    if vuln.get('url'):
                        doc.add_paragraph(f'Location: {vuln.get("url")}')
                    
                    if vuln.get('description'):
                        doc.add_paragraph(f'Description: {vuln.get("description")}')
                    
                    if vuln.get('remediation'):
                        doc.add_paragraph(f'Remediation: {vuln.get("remediation")}')
                    
                    doc.add_paragraph()  # Add spacing
            
            # Recommendations
            doc.add_heading('Recommendations', level=1)
            
            for rec in enhanced_data['recommendations']:
                doc.add_heading(f'{rec["title"]} ({rec["priority"]} Priority)', level=2)
                doc.add_paragraph(rec['description'])
                doc.add_paragraph(f'Timeline: {rec["timeline"]}')
            
            # OWASP Coverage
            doc.add_heading('OWASP Top 10 Coverage', level=1)
            
            # Create table for OWASP coverage
            owasp_table = doc.add_table(rows=1, cols=3)
            owasp_table.style = 'Table Grid'
            owasp_hdr = owasp_table.rows[0].cells
            owasp_hdr[0].text = 'Category'
            owasp_hdr[1].text = 'Name'
            owasp_hdr[2].text = 'Vulnerabilities Found'
            
            for category_id, category_data in enhanced_data['owasp_coverage'].items():
                row = owasp_table.add_row().cells
                row[0].text = category_id
                row[1].text = category_data['name']
                row[2].text = str(category_data['vulnerabilities_found'])
            
            # Save document to bytes
            from io import BytesIO
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return doc_buffer.getvalue()
            
        except ImportError:
            # Fallback: Generate structured text report
            enhanced_data = self.export_enhanced_results()
            
            text_report = f"""SECURITY ASSESSMENT REPORT
{'='*50}

Target: {self.target_url}
Generated: {enhanced_data['scan_metadata']['scan_date']}
Scanner Version: {enhanced_data['scan_metadata']['scanner_version']}

EXECUTIVE SUMMARY
{'-'*20}
Risk Level: {enhanced_data['executive_summary']['risk_level']}
Status: {enhanced_data['executive_summary']['status']}

{enhanced_data['executive_summary']['summary']}

KEY FINDINGS:
"""
            
            for finding in enhanced_data['executive_summary']['key_findings']:
                text_report += f"• {finding}\n"
            
            text_report += f"\n\nVULNERABILITY BREAKDOWN\n{'-'*25}\n"
            
            for severity, count in enhanced_data['executive_summary']['severity_distribution'].items():
                text_report += f"{severity}: {count}\n"
            
            if self.vulnerabilities:
                text_report += f"\n\nDETAILED FINDINGS\n{'-'*20}\n"
                
                for i, vuln in enumerate(self.vulnerabilities[:10], 1):
                    text_report += f"\n{i}. {vuln.get('type', 'Unknown')}\n"
                    text_report += f"   Severity: {vuln.get('severity', 'Unknown')}\n"
                    text_report += f"   CVSS: {vuln.get('cvss_score', 'N/A')}\n"
                    if vuln.get('description'):
                        text_report += f"   Description: {vuln.get('description')}\n"
            
            text_report += f"\n\nRECOMMENDATIONS\n{'-'*15}\n"
            
            for rec in enhanced_data['recommendations']:
                text_report += f"\n{rec['priority']} PRIORITY: {rec['title']}\n"
                text_report += f"Timeline: {rec['timeline']}\n"
                text_report += f"{rec['description']}\n"
            
            return text_report.encode('utf-8')
    
    def test_individual_scanners(self):
        """Test individual scanner components - legacy method for compatibility"""
        # This method exists for backward compatibility with the enhanced scan type
        additional_vulns = []
        
        # Run some additional lightweight tests
        try:
            # Test for robots.txt disclosure
            robots_response = self.session.get(f"{self.target_url}/robots.txt", timeout=5)
            if robots_response.status_code == 200 and 'disallow' in robots_response.text.lower():
                additional_vulns.append({
                    'type': 'Information Disclosure',
                    'severity': 'Low',
                    'cvss_score': 3.7,
                    'cvss_vector': 'CVSS:3.1/AV:N/AC:H/PR:N/UI:N/S:U/C:L/I:N/A:N',
                    'category': 'A05',
                    'owasp_category': 'A05:2021 - Security Misconfiguration',
                    'url': f"{self.target_url}/robots.txt",
                    'description': 'robots.txt file exposes directory structure and hidden paths',
                    'impact': 'Information leakage about application structure',
                    'remediation': 'Review robots.txt content and remove sensitive path information'
                })
        except:
            pass
        
        self.vulnerabilities.extend(additional_vulns)
        return additional_vulns